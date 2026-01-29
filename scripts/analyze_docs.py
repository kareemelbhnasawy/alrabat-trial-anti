import fitz  # PyMuPDF
from docx import Document
import os

pdf_path = "Alrabat Company Profile V4.pdf"
docx_path = "Website comments (1).docx"

def extract_pdf(path):
    print(f"\n--- EXTRACTING PDF: {path} ---")
    if not os.path.exists(path):
        print("File not found.")
        return
    try:
        doc = fitz.open(path)
        for i, page in enumerate(doc):
            print(f"\n[Page {i+1}]")
            print(page.get_text())
    except Exception as e:
        print(f"Error reading PDF: {e}")

def extract_docx(path):
    print(f"\n--- EXTRACTING DOCX: {path} ---")
    if not os.path.exists(path):
        print("File not found.")
        return
    try:
        doc = Document(path)
        for para in doc.paragraphs:
            if para.text.strip():
                print(para.text)
        
        # Also check tables if any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                print(" | ".join(row_text))
    except Exception as e:
        print(f"Error reading DOCX: {e}")

if __name__ == "__main__":
    extract_pdf(pdf_path)
    extract_docx(docx_path)
